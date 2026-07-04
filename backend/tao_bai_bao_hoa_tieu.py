from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    Usage:
    set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                          bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"})
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_run_font(run, font_name='Times New Roman', size=12, bold=False, italic=False, color=None):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(*color) if color else RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name='Times New Roman',
                  size=12, bold=False, italic=False, space_after=6, color=None, first_line_indent=0.0):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading_news(doc, text, level=1):
    """Tiêu đề bài báo"""
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(text)
        set_run_font(run, font_name='Times New Roman', size=16, bold=True, color=(0, 0, 139))
        # Thêm đường gạch chân underline cho tiêu đề
        run.font.underline = False
        # Thêm border dưới qua paragraph PPr
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000080')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(text)
        set_run_font(run, font_name='Times New Roman', size=14, bold=True, color=(0, 0, 128))
    return p

def add_lead_paragraph(doc, text):
    """Đoạn lead - thụt lề đầu dòng 1cm"""
    return add_paragraph(
        doc, text,
        bold=True,
        size=13,
        space_after=8,
        first_line_indent=0.394  # 1cm ~ 0.394 inches
    )

def add_body_paragraph(doc, text):
    """Đoạn thân bài - thụt lề đầu dòng 0.5cm"""
    return add_paragraph(
        doc, text,
        size=12,
        space_after=6,
        first_line_indent=0.197  # 0.5cm ~ 0.197 inches
    )

def add_author_dateline(doc, author, location):
    """Thêm dòng tác giả - địa điểm"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f'{location} | {author}')
    set_run_font(run, font_name='Times New Roman', size=11, italic=True, color=(80, 80, 80))
    return p

def add_horizontal_line(doc):
    """Đường kẻ ngang"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0C0C0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)

def add_sub_heading(doc, text):
    """Tiểu mục"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    set_run_font(run, font_name='Times New Roman', size=13, bold=True, color=(0, 0, 128))
    return p

def add_caption(doc, text):
    """Chú thích ảnh hoặc bảng"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    set_run_font(run, font_name='Times New Roman', size=11, italic=True, color=(60, 60, 60))
    return p

# ==========================================
# TẠO BÀI BÁO TIN TỨC
# ===========================================
doc = Document()

# --- Cài đặt lề trang ---
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# --- HEADER: Tên báo / chuyên mục ---
add_paragraph(doc, 'BÁO GIAO THÔNG VẬN TẢI', alignment=WD_ALIGN_PARAGRAPH.CENTER,
              font_name='Arial', size=11, bold=False, color=(100, 100, 100), space_after=2)
add_paragraph(doc, 'CHUYÊN MỤC: HÀNG HẢI', alignment=WD_ALIGN_PARAGRAPH.CENTER,
              font_name='Arial', size=11, bold=False, color=(0, 0, 128), space_after=4)
add_horizontal_line(doc)

# --- TIÊU ĐỀ CHÍNH ---
add_heading_news(doc, 'NGHỀ HOA TIÊU HÀNG HẢI: GIỮ VỊ TRÍ BẢO VỆ AN TOÀN VẬN CHUYỂN BIỂN', level=1)

# Dòng tác giả - địa điểm - thời gian
add_author_dateline(doc, author='PV Báo GTVT', location='Hà Nội')

# Tóm tắt bài viết (lead / lead paragraph)
lead = (
    'Cả nước hiện có hàng trăm hoa tiêu chuyên nghiệp đang làm nhiệm vụ dẫn tàu vào các cảng biển lớn. '
    'Nghề hoa tiêu không chỉ đảm bảo an toàn hàng hải mà còn góp phần nâng cao hiệu quả vận tải biển '
    'và tối ưu chi phí khai thác cảng, bất chấp những thách thức từ biến đổi khí hậu, tự động hóa và '
    'áp lực cạnh tranh về giá dịch vụ.'
)
add_lead_paragraph(doc, lead)

# ===== NỘI DUNG CHÍNH =====
add_horizontal_line(doc)

body1 = (
    'Tại Việt Nam, nghề hoa tiêu hàng hải đã phát triển qua nhiều giai đoạn, từ các hoa tiêu địa phương '
    'tự nguyện hỗ trợ tàu vào cảng ven biển đến một mô hình chuyên nghiệp hóa, có tổ chức quản lý trên '
    'phạm vi toàn quốc. Hiện nay, cả nước có 3 công ty hoa tiêu hàng hải trực thuộc Công ty TNHH MTV '
    'Hoa tiêu hàng hải Việt Nam, phân chia theo miền: miền Bắc, miền Trung và miền Nam. Các công ty '
    'này đều hoạt động theo mô hình độc quyền với sự tham gia của hoa tiêu là các công dân Việt Nam, '
    'có chứng chỉ chuyên môn theo quy định của Bộ Giao thông vận tải.'
)
add_body_paragraph(doc, body1)

body2 = (
    'Hoa tiêu là người có chứng chỉ chuyên môn được cơ quan nhà nước có thẩm quyền cấp, có kinh nghiệm '
    'và nắm vững đặc thù luồng, vùng nước, điều kiện khí tượng, hải lưu của cảng biển để trực tiếp dẫn '
    'đường cho tàu vào cảng, rời cảng hoặc di chuyển trong các khu vực có khó khăn về an toàn hàng hải. '
    'Kể từ năm 2021, Bộ Giao thông vận tải đã ban hành các văn bản quy định chặt chẽ về tiêu chuẩn '
    'đào tạo, cấp, thu hồi chứng chỉ và giới hạn phạm vi hoạt động của hoa tiêu, bao gồm phân hạng từ '
    'Hạng Ba, Hạng Nhì, Hạng Nhất đến Ngoại hạng, tương ứng với tàu có dung tích từ 5.000GT đến không '
    'giới hạn.'
)
add_body_paragraph(doc, body2)

body3 = (
    'Nghề hoa tiêu được ví như "con mắt" của cảng biển. Theo thống kê của các công ty hoa tiêu, số '
    'lượng tàu cần hoa tiêu hướng dẫn mỗi năm tại các cảng lớn như Cái Mép, Cát Lái, Sài Gòn, Hải '
    'Phòng lên đến hàng nghìn lượt, trong đó phần lớn là tàu có trọng tải lớn, khó thao tác do dung '
    'tích vượt quá 10.000GT. Một hoa tiêu Ngoại hạng có thể dẫn tàu có dung tích lên đến hàng chục '
    'nghìn tấn, cần nhiều giờ để đưa tàu qua những khúc luồng hẹp, nơi có nhiều tàu khác đang cập '
    'cảng, xuôi ngược.'
)
add_body_paragraph(doc, body3)

# Tiểu mục 1
add_sub_heading(doc, 'Ý nghĩa của hoa tiêu đối với an toàn và hiệu quả vận tải biển')

body4 = (
    'Việc hoa tiêu trực tiếp dẫn tàu giúp chủ tàu, cảng vụ và các đơn vị liên quan chủ động dự báo '
    'rủi ro, tính toán thời gian neo đậu, giảm thiểu chi phí chờ đợi và hạn chế tối đa tai nạn hàng '
    'hải. Đặc biệt tại các cảng có luồng cảng hẹp, đông tàu hoặc có điều kiện tự nhiên phức tạp như '
    'cảng Sài Gòn, cảng Hải Phòng, cảng Quy Nhơn, vai trò của hoa tiêu càng trở nên thiết yếu.'
)
add_body_paragraph(doc, body4)

body5 = (
    'Bên cạnh đó, hoa tiêu còn đóng vai trò tư vấn về điều kiện an toàn, hướng dẫn thuyền trưởng '
    'nắm rõ quy định về luồng, khu vực neo đậu, thời gian cao điểm và các hạn chế khi lưu thông '
    'trong cảng. Nhờ đó, các cảng biển Việt Nam duy trì được tỷ lệ tai nạn thấp và uy tín trong mắt '
    'các hãng tàu quốc tế.'
)
add_body_paragraph(doc, body5)

# Tiểu mục 2
add_sub_heading(doc, 'Chuẩn hóa công tác hoa tiêu: Thông tư mới có hiệu lực từ 30/12/2025')

body6 = (
    'Ngày 18/11/2025, Bộ Giao thông vận tải đã ban hành Thông tư quy định chi tiết về tiêu chuẩn '
    'đào tạo, cấp và thu hồi chứng chỉ chuyên môn, cùng giấy chứng nhận vùng hoạt động hoa tiêu hàng '
    'hải, thay thế hai thông tư cũ số 27/2016/TT-BGTVT và số 54/2023/TT-BGTVT. Thông tư này chính '
    'thức có hiệu lực từ ngày 30/12/2025 và áp dụng đến tất cả hoa tiêu hàng hải, thuyền trưởng tự '
    'dẫn, các tổ chức hoa tiêu hàng hải, chủ tàu và các cơ quan, tổ chức liên quan.'
)
add_body_paragraph(doc, body6)

body7 = (
    'Theo đó, chứng chỉ khả năng chuyên môn của hoa tiêu được chia thành 4 hạng: Hạng Ba (dẫn tàu '
    'tối đa 5.000GT và 115m), Hạng Nhì (tối đa 10.000GT và 145m), Hạng Nhất (tối đa 20.000GT và '
    '175m) và Ngoại hạng (không giới hạn). Thời hạn sử dụng chứng chỉ là 5 năm kể từ ngày cấp. Để '
    'nâng hạng, hoa tiêu phải tích lũy đủ số lượt dẫn tàu an toàn và thời gian công tác theo quy '
    'định, qua đó khuyến khích người hành nghề không ngừng nâng cao năng lực.'
)
add_body_paragraph(doc, body7)

body8 = (
    'Thông tư cũng quy định các hình thức xử lý nghiêm đối với trường hợp hoa tiêu có lỗi dẫn đến '
    'tai nạn hàng hải, bao gồm kéo dài thời gian và bổ sung số lượt dẫn tàu an toàn để được nâng '
    'hạng. Quy định này nhằm củng cố ý thức trách nhiệm và công tâm trong nghề, đồng thời khẳng định '
    'vị thế của nghề hoa tiêu như một ngành nghề đẳng cấp, chuyên sâu.'
)
add_body_paragraph(doc, body8)

# Tiểu mục 3
add_sub_heading(doc, 'Giảm giá dịch vụ hoa tiêu: hỗ trợ tàu Việt Nam, thúc đẩy phục hồi kinh tế biển')

body9 = (
    'Cuối năm 2024, Bộ Giao thông vận tải đã ban hành Quyết định số 814/QĐ-BGTVT quy định giá '
    'tối đa dịch vụ hoa tiêu hàng hải tại các cảng biển Việt Nam. Theo đó, giá dịch vụ hoa tiêu được '
    'điều chỉnh theo từng hạng tàu và từng khu vực, nhằm giảm chi phí vận hành cho chủ tàu trong '
    'bối cảnh giá nhiên liệu và chi phí logistics toàn cầu tăng cao.'
)
add_body_paragraph(doc, body9)

body10 = (
    'Gần đây, các tổ chức hoa tiêu hàng hải trong cả nước đã đi đến thống nhất giảm giá dịch vụ '
    'hoa tiêu tối thiểu 10% đối với tàu mang quốc tịch Việt Nam, thể hiện chính sách hỗ trợ doanh '
    'nghiệp vận tải biển nội địa, thúc đẩy phục hồi và phát triển đội tàu biển quốc gia. Đồng thời, '
    'các công ty hoa tiêu hàng hải miền Bắc, miền Nam đều hoàn thành mục tiêu tái cơ cấu thành công, '
    'tăng mạnh sản lượng và lợi nhuận năm 2025, đề ra kế hoạch bứt phá năm 2026.'
)
add_body_paragraph(doc, body10)

body11 = (
    'Việc tái cơ cấu các công ty hoa tiêu không chỉ giúp công ty hoạt động hiệu quả hơn mà còn nâng '
    'cao chất lượng dịch vụ, đẩy nhanh tốc độ xử lý tàu vào cảng và rút ngắn thời gian chờ đợi. '
    'Những kết quả này góp phần nâng cao năng lực cạnh tranh của ngành cảng biển Việt Nam trong bối '
    'cảnh hội nhập quốc tế.'
)
add_body_paragraph(doc, body11)

# Tiểu mục 4
add_sub_heading(doc, 'Thách thức và định hướng phát triển trong tương lai')

body12 = (
    'Mặc dù có nhiều kết quả đáng ghi nhận, nghề hoa tiêu vẫn đối mặt với không ít thách thức. '
    'Một trong những khó khăn lớn nhất là thiếu hụt nhân lực hoa tiêu chất lượng cao, đặc biệt ở '
    'hạng Ngoại hạng, do yêu cầu cao về kinh nghiệm, trình độ ngoại ngữ và khả năng xử lý tình huống '
    'phức tạp. Tỷ lệ hoa tiêu đủ điều kiện nâng hạng còn thấp so với nhu cầu thực tế của các cảng '
    'biển lớn.'
)
add_body_paragraph(doc, body12)

body13 = (
    'Bên cạnh đó, biến đổi khí hậu đang làm thay đổi đáng kể điều kiện hải lưu, yếu tố thời tiết, '
    'đòi hỏi hoa tiêu phải nắm vững kiến thức mới để đảm bảo an toàn. Sự phát triển của công nghệ '
    'tự động hóa, hệ thống định vị và mô phỏng dẫn tàu cũng tác động đến công việc truyền thống của '
    'hoa tiêu, buộc ngành phải đổi mới phương pháp đào tạo và ứng dụng công nghệ vào hoạt động.'
)
add_body_paragraph(doc, body13)

body14 = (
    'Trong thời gian tới, các chuyên gia khuyến nghị Nhà nước cần tiếp tục hoàn thiện khung pháp lý, '
    'đầu tư cơ sở vật chất, đào tạo nguồn nhân lực có chất lượng và mở rộng hợp tác quốc tế để '
    'nâng cao năng lực nghề hoa tiêu Việt Nam. Đây là yếu tố quan trọng góp phần đưa Việt Nam trở '
    'thành trung tâm vận tải biển khu vực và thế giới.'
)
add_body_paragraph(doc, body14)

# Đoạn kết
kết_luận = (
    'Nghề hoa tiêu hàng hải vẫn giữ vị trí đặc biệt trong hệ thống giao thông vận tải biển Việt Nam. '
    'Với sự chuẩn hóa ngày càng cao, việc đầu tư phát triển nguồn nhân lực chất lượng và áp dụng '
    'công nghệ hiện đại sẽ góp phần giữ vững an toàn hàng hải và nâng cao vị thế ngành cảng biển '
    'nước nhà trên bản đồ quốc tế.'
)
add_body_paragraph(doc, kết_luận)

# Đường kẻ cuối
add_horizontal_line(doc)

# --- GÓC BÌA / THÔNG TIN XUẤT BẢN ---
add_paragraph(doc, 'Tác giả: PV Báo Giao thông Vận tải', alignment=WD_ALIGN_PARAGRAPH.LEFT,
              font_name='Times New Roman', size=11, color=(80, 80, 80), space_after=2)
add_paragraph(doc, 'Nguồn tham khảo: Báo Nhân Dân, Bộ Giao thông vận tải, VMSA, VIMC, các công ty hoa tiêu hàng hải',
              alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name='Times New Roman', size=11, color=(80, 80, 80), space_after=2)
ngay = datetime.now().strftime('%d/%m/%Y')
add_paragraph(doc, f'Xuất bản: {ngay} | Thể loại: Tin tức hàng hải',
              alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name='Times New Roman', size=11, color=(80, 80, 80), space_after=0)

# Lưu file
output_dir = os.path.dirname(__file__) if '__file__' in dir() else '.'
output_path = os.path.join(output_dir, 'Bai_bao_tin_tuc_Hoa_tieu.docx')
doc.save(output_path)
print(f"Đã lưu bài báo tại: {output_path}")
